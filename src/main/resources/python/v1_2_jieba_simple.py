from docx import Document
import jieba
import jieba.posseg as posseg
import jieba.analyse as analyse
import re
import numpy as np
import pandas as pd

#列表循环读取
policy_warehouse = [] # 以字典存储整个文章的每个整句话，字典属性：句子本身、关键字、第几章、第几条
policy_userinput = [] # 以字典存储用户输入的文档中每句话的关键字，字典属性：句子本身、关键字、用户输入的第几句

def document_analysis():
    #打开word文档
    document = Document("D:\全部资料\研究生学前学习\政策NLP demo\检测文件-甘井子区经济工作奖励暂行办法.docx")
    document_input = Document("D:\全部资料\研究生学前学习\政策NLP demo\规则-上级文档.docx")

    #获取所有段落
    all_paragraphs = document.paragraphs # 这是个列表

    #读取政策库所有政策
    zhang = '第一章'
    for paragraph in all_paragraphs:
        txt = paragraph.text
        pattern = r'/|\'|`|\?|"|!|。|·|！| |…|'
        result_list = re.split(pattern, txt)  # 每一段都按整句来分句
        tiao = '第一条'
        for res in result_list:
            if(res != ""): # 筛掉空字符串
                tr_result = analyse.textrank(res, topK=10) # 分析关键字

                pattern1 = re.compile(r'第\w+章') # 检测这句话是否是"第x章"
                result1 = pattern1.findall(res)
                if(len(result1)!=0):
                    zhang = result1[0]

                pattern2 = re.compile(r'第\w+条') # 检测这句话是否是"第x条"
                result2 = pattern2.findall(res)
                if(len(result2)!=0):
                    tiao = result2[0]
                    
                policy_warehouse.append({'sentence': res, 'keywords': tr_result, 'zhang': zhang, 'tiao': tiao})

    all_paragraphs = document_input.paragraphs

    #读取用户输入文档的查询语句
    for index_para, paragraph in enumerate(all_paragraphs):
        txt = paragraph.text
        pattern = r'/|\'|`|\?|"|!|。|·|！| |…|'
        result_list = re.split(pattern, txt)  # 每一段都按整句来分句
        for index_sen, res in enumerate(result_list):
            tr_result = analyse.textrank(res, topK=10) # 分析关键字
            if(len(tr_result)!=0 and index_para!=0): # 跳过第一句的“大连市文件”
                tr_result = analyse.textrank(res, topK=10)
                policy_userinput.append({'sentence': res, 'keywords': tr_result, 'number': index_para})
    
    '''为匹配列表处理匹配度'''
    match_list = keyword_match() # 遍历两个文档以获取匹配序列
    def key_return(e): # 定义一个返回'match_count'键的函数
        return e['match_count']
    match_list.sort(key=key_return, reverse=True) # 倒序(由高到低)，且以字典的match_count的值来排序，排不排序应该都行，排完看着舒服
    
    count_list = [] # 创建一个只存储match_list中match_count数据的一维数组
    for m in match_list: # 我这里使用for循环创建的，有没有办法用一个函数就能取出来这一列？
        count_list.append(m['match_count'])
    count_list = np.array(count_list) # list转array
    count_list = pd.cut(count_list, bins=3, labels=False) # 数据自动分箱，分箱(bins=3个箱子)操作后，返回值为0,1,2的数组，表示"低,中,高"
    print("分箱结果:"+str(count_list)) # 看一下分箱结果吧！# pd.cut函数设置参数retbins=True，函数就可以返回分箱的边界值
    for index, m in enumerate(match_list):
        m['matching_degree'] = count_list[index] # 为match_list增加"匹配度"键值对，2为高，1为中，0为低
    return match_list

# 检测特定索引i,j两句话的关键字匹配数
def two_sentence_match(index1, index2): 
    count = 0
    for keyword1 in policy_warehouse[index1]['keywords']:
        for keyword2 in policy_userinput[index2]['keywords']:
            if keyword1 == keyword2:
                count += 1
    return count

# 遍历政策库和用户输入所有句子
def keyword_match():
    match_list = []
    for i in range(len(policy_warehouse)):
        for j in range(len(policy_userinput)):
            count = two_sentence_match(i, j)
            if(count > 2): # 把关键字匹配门槛设置为2，高于2的才有资格评判其匹配度是高/中/低
                match_list.append({'warehouse': policy_warehouse[i], 'userinput': policy_userinput[j], 'match_count': count})
    return match_list

document_analysis()