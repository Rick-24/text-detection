from docx import Document
import jieba
import jieba.posseg as posseg
import jieba.analyse as analyse
import re
import numpy as np
import pandas as pd

import json
import urllib.request
import threading
import socket


#match_list:存储匹配后 检测文档的句子本身、关键字、第几章、第几条；匹配规则文档的句子本身、关键字、第几条规则
'''
match_list格式:
[{ 'warehouse': #对应（左）检测文档--第五章第四条
   {   'sentence': '对年度实现规模以上工业总产值在全区排前五名的街道，分别奖励50万元、40万元、30万元、20万元、10万元',
       'keywords': ['工业', '总产值', '规模', '实现', '全区', '街道', '年度', '奖励'],
       'zhang': '第五章',
       'tiao': '第四条',
       'duan':3, #在原文中的第几段
       'ju':1 #在这一段的第几句
   },
   'userinput': #对应（右）规则文档--第二条规则
   {   'sentence': '每年工业总产值在全区排前二十名的街道，分别奖励50万元、30万元、20万元',
       'keywords': ['全区', '街道', '总产值', '工业', '奖励'],
       'number': 2,
       'duan':2, #在原文中的第几段
       'ju':1 #在这一段的第几句
   },
   'match_count': 5,    #相同关键词个数
   'matching_degree': 2 #匹配度等级为2级
 },...]
'''

def document_analysis(input_doc_name):
    #打开word文档
    document = Document("D:\全部资料\研究生学前学习\政策NLP demo\检测文件-甘井子区经济工作奖励暂行办法.docx")
    document_input = Document(input_doc_name)
    
    
    policy_warehouse = [] # 以字典存储整个文章的每个整句话，字典属性：句子本身、关键字、第几章、第几条
    policy_userinput = [] # 以字典存储用户输入的文档中每句话的关键字，字典属性：句子本身、关键字、用户输入的第几条规则
    
    #获取所有段落
    all_paragraphs = document.paragraphs # 这是个列表
    #读取政策库所有政策
    zhang = '第一章'
    print("原文档的段落数：",len(all_paragraphs))
    for index_para,paragraph in enumerate(all_paragraphs):
        txt = paragraph.text
        pattern = r'/|\'|`|\?|"|!|。|·|！| |…|'
        result_list = re.split(pattern, txt)  # 每一段都按整句来分句
        #print("duan%s"%(index_para+1))
        tiao = '第一条'
        for ju,res in enumerate(result_list):
            #print("ju:",ju+1)
            if(res != ""): # 筛掉空字符串
                tr_result = analyse.textrank(res, topK=10) # 分析关键字
                pattern1 = re.compile(r'第\w+章') # 检测这句话是否是"第x章"
                result1 = pattern1.findall(res)
                if(len(result1)!=0):
                    zhang = result1[0]

                pattern2 = re.compile(r'第\w+条') # 检测这句话是否是"第x条"
                result2 = pattern2.findall(res)
                if(len(result2)!=0):
                    tiao = result2[0]

                policy_warehouse.append({'sentence': res, 'keywords': tr_result, 'zhang': zhang, 'tiao': tiao,'duan': str(index_para+1),'ju':str(ju+1)})

    all_paragraphs = document_input.paragraphs
    #读取用户输入文档的查询语句
    for index_para, paragraph in enumerate(all_paragraphs):
        txt = paragraph.text
        pattern = r'/|\'|`|\?|"|!|。|·|！| |…|'
        result_list = re.split(pattern, txt)  # 每一段都按整句来分句
        for index_sen, res in enumerate(result_list):
            tr_result = analyse.textrank(res, topK=10) # 分析关键字
            if(len(tr_result)!=0 and index_para!=0): # 跳过第一句的“大连市文件”
                #tr_result = analyse.textrank(res, topK=10)#分析关键字

                policy_userinput.append({'sentence': res, 'keywords': tr_result, 'number': str(index_para),'duan':str(index_para+1),'ju':str(index_sen+1)})
    
    '''为匹配列表处理匹配度'''
    match_list = keyword_match(policy_warehouse, policy_userinput) # 遍历两个文档以获取匹配序列
    def key_return(e): # 定义一个返回'match_count'键的函数
        return e['match_count']
    match_list.sort(key=key_return, reverse=True) # 倒序(由高到低)，且以字典的match_count的值来排序，排不排序应该都行，排完看着舒服
    
    count_list = [] # 创建一个只存储match_list中match_count数据的一维数组
    for m in match_list: # 我这里使用for循环创建的，有没有办法用一个函数就能取出来这一列？
        count_list.append(m['match_count'])
    count_list = np.array(count_list) # list转array
    count_list = pd.cut(count_list, bins=3, labels=False) # 数据自动分箱，分箱(bins=3个箱子)操作后，返回值为0,1,2的数组，表示"低,中,高"
    print("分箱结果:"+str(count_list)) # 看一下分箱结果吧！# pd.cut函数设置参数retbins=True，函数就可以返回分箱的边界值
    for index, m in enumerate(match_list):
        m['match_count'] = str(m['match_count'])
        m['matching_degree'] = str(count_list[index]) # 为match_list增加"匹配度"键值对，2为高，1为中，0为低

    return match_list

# 检测特定索引i,j两句话的关键字匹配数
def two_sentence_match(index1, index2, policy_warehouse, policy_userinput): 
    count = 0
    for keyword1 in policy_warehouse[index1]['keywords']:
        for keyword2 in policy_userinput[index2]['keywords']:
            if keyword1 == keyword2:
                count += 1
    return count

# 遍历政策库和用户输入所有句子（遍历两个文档以获取匹配序列）
def keyword_match(policy_warehouse, policy_userinput):
    match_list = []
    for i in range(len(policy_warehouse)):
        for j in range(len(policy_userinput)):
            count = two_sentence_match(i, j, policy_warehouse, policy_userinput)#返回两个句子 相同关键字的数量
            if(count > 2): # 把关键字匹配门槛设置为2，高于2的才有资格评判其匹配度是高/中/低
                match_list.append({'warehouse': policy_warehouse[i], 'userinput': policy_userinput[j], 'match_count': count})
    return match_list

encoding = 'utf-8'
BUFSIZE = 1024

# a read thread, read data from remote
class Reader(threading.Thread):
    def __init__(self, client):
        threading.Thread.__init__(self)
        self.client = client
        
    def run(self):
        while True:
            data = self.client.recv(BUFSIZE)
            if(data):
                input = bytes.decode(data, encoding)
                for i, c in enumerate(input):
                    if c=='{':
                        input_doc_name = input[i:]
                        break
                input_doc_name = json.loads(input_doc_name)
                input_doc_name = input_doc_name['doc_name']
                
                match_list = document_analysis(input_doc_name)
                
                headers = {'Content-Type': 'application/json'}
                #print(json.dumps(data).encode()) # dict先转json，再转二进制，才能进行传输
                request = urllib.request.Request(url='http://localhost:2453/', headers=headers, data=json.dumps(match_list).encode())
                response = urllib.request.urlopen(request)
                
                #print(match_list)
            else:
                break
        print("close:", self.client.getpeername())
        
    def readline(self):
        rec = self.inputs.readline()
        if rec:
            string = bytes.decode(rec, encoding)
            if len(string)>2:
                string = string[0:-2]
            else:
                string = ' '
        else:
            string = False
        return string

# a listen thread, listen remote connect
# when a remote machine request to connect, it will create a read thread to handle
class Listener(threading.Thread):
    def __init__(self, port):
        threading.Thread.__init__(self)
        self.port = port
        self.sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        self.sock.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
        self.sock.bind(("0.0.0.0", port))
        self.sock.listen(0)
    def run(self):
        print("listener started")
        while True:
            client, cltadd = self.sock.accept()
            Reader(client).start()
            cltadd = cltadd
            #print("accept a connect")

lst  = Listener(2452)   # create a listen thread
lst.start() # then start

# input_doc_name = "C:/Users/DerBo/Desktop/input1.docx"
# document_analysis(input_doc_name)
